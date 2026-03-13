"""
Service de génération de documents Word (DOCX) professionnels
Institut Gabriel Rita - Service Rapports
"""
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Couleurs IGR
COLOR_PRIMARY = RGBColor(11, 75, 131)   # Bleu principal
COLOR_ACCENT  = RGBColor(230, 166, 71)  # Or/Orange accent
COLOR_WHITE   = RGBColor(255, 255, 255)
COLOR_GRAY    = RGBColor(107, 114, 128)
COLOR_LIGHT   = RGBColor(243, 244, 246)
COLOR_RED     = RGBColor(220, 38, 38)
COLOR_GREEN   = RGBColor(16, 185, 129)


def _set_cell_bg(cell, hex_color: str):
    """Définir la couleur de fond d'une cellule"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Ajouter des bordures à une cellule"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def _add_header(doc: Document, title: str, subtitle: str = ''):
    """Ajouter un en-tête professionnel IGR"""
    # Table header: logo zone + titre + date
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Colonne 1: École
    c0 = table.cell(0, 0)
    _set_cell_bg(c0, '0B4B83')
    c0.width = Cm(5)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('🏫  ISGR')
    run.bold = True
    run.font.color.rgb = COLOR_WHITE
    run.font.size = Pt(14)
    p2 = c0.add_paragraph('Institut Gabriel Rita')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0] if p2.runs else p2.add_run('Institut Gabriel Rita')
    r2.font.color.rgb = COLOR_ACCENT
    r2.font.size = Pt(9)

    # Colonne 2: Titre central
    c1 = table.cell(0, 1)
    _set_cell_bg(c1, '1B527E')
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p3 = c1.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(title.upper())
    r3.bold = True
    r3.font.color.rgb = COLOR_WHITE
    r3.font.size = Pt(13)
    if subtitle:
        p4 = c1.add_paragraph(subtitle)
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.runs[0] if p4.runs else p4.add_run(subtitle)
        r4.font.color.rgb = COLOR_ACCENT
        r4.font.size = Pt(9)

    # Colonne 3: Date
    c2 = table.cell(0, 2)
    _set_cell_bg(c2, '0B4B83')
    c2.width = Cm(4)
    p5 = c2.paragraphs[0]
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run('Année scolaire')
    r5.font.color.rgb = COLOR_ACCENT
    r5.font.size = Pt(8)
    p6 = c2.add_paragraph('2025 - 2026')
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.runs[0] if p6.runs else p6.add_run('2025-2026')
    r6.bold = True
    r6.font.color.rgb = COLOR_WHITE
    r6.font.size = Pt(11)
    p7 = c2.add_paragraph(datetime.now().strftime('%d/%m/%Y'))
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r7 = p7.runs[0] if p7.runs else p7.add_run('')
    r7.font.color.rgb = COLOR_WHITE
    r7.font.size = Pt(8)

    doc.add_paragraph()  # Espace


def _add_footer(doc: Document):
    """Ajouter un pied de page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Document généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")} — Institut Gabriel Rita © {datetime.now().year}')
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True


def _styled_table_header(table, headers: list, col_widths: list = None):
    """Créer une ligne d'en-tête stylée pour un tableau"""
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, '0B4B83')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = COLOR_WHITE
        run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            hdr_row.cells[i].width = Cm(w)


class DocxService:

    @staticmethod
    def generate_schedule(data: dict) -> bytes:
        """Génère l'emploi du temps d'une classe en DOCX"""
        doc = Document()
        # Marges serrées pour A4 paysage
        section = doc.sections[0]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.orientation   = 1  # Paysage

        _add_header(doc, f"Emploi du Temps", f"Classe : {data['class_name']}")

        days = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi']
        time_slots = data.get('time_slots', [])
        schedule = data.get('schedule', {})

        # Créer tableau: colonne Horaire + 6 jours
        table = doc.add_table(rows=1 + len(time_slots), cols=7)
        table.style = 'Table Grid'

        # En-têtes
        headers = ['Horaire'] + days
        _styled_table_header(table, headers)

        # Données
        for row_idx, time_slot in enumerate(time_slots):
            row = table.rows[row_idx + 1]
            # Heure
            tc = row.cells[0]
            _set_cell_bg(tc, 'E6A647')
            p = tc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(time_slot)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(17, 24, 39)

            # Jours
            for col_idx, day in enumerate(days):
                cell = row.cells[col_idx + 1]
                slot = schedule.get(day, {}).get(time_slot)
                if slot:
                    # Fond alternant
                    _set_cell_bg(cell, 'EFF6FF')
                    p = cell.paragraphs[0]
                    r_subj = p.add_run(slot.get('subject', ''))
                    r_subj.bold = True
                    r_subj.font.size = Pt(8)
                    r_subj.font.color.rgb = COLOR_PRIMARY

                    p2 = cell.add_paragraph(f"👤 {slot.get('teacher_name', '')}")
                    if p2.runs: p2.runs[0].font.size = Pt(8.5)
                else:
                    _set_cell_bg(cell, 'F9FAFB')

        doc.add_paragraph()
        _add_footer(doc)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_student_report(data: dict) -> bytes:
        """Génère le relevé de compte d'un étudiant en DOCX"""
        doc = Document()
        section = doc.sections[0]
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

        student = data['student']
        _add_header(doc, "Relevé de Compte Étudiant",
                    f"{student['firstName']} {student['lastName']} — {student['matricule']}")

        # Infos étudiant
        p = doc.add_paragraph()
        p.add_run('INFORMATIONS ÉTUDIANT').bold = True
        p.runs[-1].font.color.rgb = COLOR_PRIMARY
        p.runs[-1].font.size = Pt(11)

        info_table = doc.add_table(rows=3, cols=4)
        info_table.style = 'Table Grid'
        infos = [
            ('Nom', student['lastName']),
            ('Prénom', student['firstName']),
            ('Matricule', student['matricule']),
            ('Classe', student['class_name']),
            ('Frais dus', f"{student['totalFeesDue']:,.0f} FCFA"),
            ('Frais payés', f"{student['totalFeesPaid']:,.0f} FCFA"),
        ]
        for i, (label, val) in enumerate(infos):
            row = i // 2
            col_base = (i % 2) * 2
            lbl_cell = info_table.cell(row, col_base)
            _set_cell_bg(lbl_cell, 'EFF6FF')
            lbl_cell.paragraphs[0].add_run(label).bold = True
            info_table.cell(row, col_base + 1).paragraphs[0].add_run(val)

        # Solde
        doc.add_paragraph()
        balance = student['balance']
        bal_color = '16A34A' if balance >= 0 else 'DC2626'
        bal_table = doc.add_table(rows=1, cols=1)
        bal_table.style = 'Table Grid'
        bc = bal_table.cell(0, 0)
        _set_cell_bg(bc, bal_color)
        bp = bc.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run(f"SOLDE : {balance:+,.0f} FCFA")
        br.bold = True
        br.font.color.rgb = COLOR_WHITE
        br.font.size = Pt(14)

        # Historique paiements
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.add_run('HISTORIQUE DES PAIEMENTS').bold = True
        p2.runs[-1].font.color.rgb = COLOR_PRIMARY

        payments = data.get('payments', [])
        if payments:
            tbl = doc.add_table(rows=1 + len(payments), cols=4)
            tbl.style = 'Table Grid'
            _styled_table_header(tbl, ['Date', 'Référence', 'Mode', 'Montant (FCFA)'])
            for i, pay in enumerate(payments):
                row = tbl.rows[i + 1]
                date_str = pay['paymentDate'].strftime('%d/%m/%Y') if hasattr(pay['paymentDate'], 'strftime') else str(pay['paymentDate'])[:10]
                row.cells[0].paragraphs[0].add_run(date_str).font.size = Pt(9)
                row.cells[1].paragraphs[0].add_run(pay.get('reference', '-')).font.size = Pt(9)
                row.cells[2].paragraphs[0].add_run(pay.get('method', '-')).font.size = Pt(9)
                r = row.cells[3].paragraphs[0].add_run(f"{pay['amount']:,.0f}")
                r.font.size = Pt(9)
                r.bold = True
                if i % 2 == 0:
                    _set_cell_bg(row.cells[3], 'F0FDF4')
        else:
            doc.add_paragraph('Aucun paiement enregistré.').italic = True

        _add_footer(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_global_school(data: dict) -> bytes:
        """Génère le rapport global de l'école en DOCX"""
        doc = Document()
        section = doc.sections[0]
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

        stats = data['statistics']
        _add_header(doc, "Rapport Global de l'École", f"Année scolaire {data['school_info']['year']}")

        # KPI boxes
        p = doc.add_paragraph()
        r = p.add_run('INDICATEURS CLÉS')
        r.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        r.font.size = Pt(11)

        kpi_table = doc.add_table(rows=2, cols=3)
        kpi_table.style = 'Table Grid'
        kpis = [
            ('Élèves inscrits', stats['total_students'], '0B4B83'),
            ('Élèves actifs', stats['active_students'], '059669'),
            ('Personnel', stats['total_staff'], '7C3AED'),
            ('Classes actives', stats['total_classes'], 'D97706'),
            ('Recettes totales', f"{stats['total_revenue']:,.0f} FCFA", '0B4B83'),
            ('', '', 'FFFFFF'),
        ]
        for i, (label, val, color) in enumerate(kpis):
            row_i = i // 3
            col_i = i % 3
            cell = kpi_table.cell(row_i, col_i)
            _set_cell_bg(cell, color)
            if label:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                r.bold = True
                r.font.size = Pt(16)
                r.font.color.rgb = COLOR_WHITE
                p2 = cell.add_paragraph(label)
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r2 = p2.runs[0] if p2.runs else p2.add_run(label)
                r2.font.size = Pt(8)
                r2.font.color.rgb = RGBColor(209, 213, 219)

        # Liste des étudiants
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"LISTE DES ÉTUDIANTS ({len(data['students'])} élèves)")
        r2.bold = True
        r2.font.color.rgb = COLOR_PRIMARY

        students = data['students']
        if students:
            tbl = doc.add_table(rows=1 + len(students), cols=5)
            tbl.style = 'Table Grid'
            _styled_table_header(tbl, ['#', 'Nom', 'Prénom', 'Classe', 'Solde (FCFA)'])
            for i, s in enumerate(students):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(str(i + 1)).font.size = Pt(8)
                row.cells[1].paragraphs[0].add_run(s['lastName']).font.size = Pt(8)
                row.cells[2].paragraphs[0].add_run(s['firstName']).font.size = Pt(8)
                row.cells[3].paragraphs[0].add_run(s['className']).font.size = Pt(8)
                bal = s['balance']
                r = row.cells[4].paragraphs[0].add_run(f"{bal:,.0f}")
                r.font.size = Pt(8)
                r.font.color.rgb = COLOR_GREEN if bal >= 0 else COLOR_RED
                if i % 2 == 0:
                    for c in row.cells:
                        _set_cell_bg(c, 'F9FAFB')

        _add_footer(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_late_payments(data: dict) -> bytes:
        """Génère le rapport des paiements en retard en DOCX"""
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        late = data.get('late_payments', [])
        _add_header(doc, "Rapport — Paiements en Retard",
                    f"Année {data['school_year']} — {len(late)} élève(s) concerné(s)")

        # Total
        total = sum(p['balance'] for p in late)
        summ = doc.add_table(rows=1, cols=2)
        summ.style = 'Table Grid'
        c0 = summ.cell(0, 0)
        _set_cell_bg(c0, 'DC2626')
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(f"Nombre d'élèves en retard : {len(late)}")
        r0.bold = True
        r0.font.color.rgb = COLOR_WHITE
        c1 = summ.cell(0, 1)
        _set_cell_bg(c1, 'B91C1C')
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(f"Total impayé : {total:,.0f} FCFA")
        r1.bold = True
        r1.font.color.rgb = COLOR_WHITE

        doc.add_paragraph()
        if late:
            tbl = doc.add_table(rows=1 + len(late), cols=5)
            tbl.style = 'Table Grid'
            _styled_table_header(tbl, ['#', 'Nom', 'Prénom', 'Classe', 'Montant dû (FCFA)'])
            for i, s in enumerate(late):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(str(i + 1)).font.size = Pt(8)
                row.cells[1].paragraphs[0].add_run(s['lastName']).font.size = Pt(8)
                row.cells[2].paragraphs[0].add_run(s['firstName']).font.size = Pt(8)
                row.cells[3].paragraphs[0].add_run(s['className']).font.size = Pt(8)
                r = row.cells[4].paragraphs[0].add_run(f"{s['balance']:,.0f}")
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = COLOR_RED
                if i % 2 == 0:
                    for c in row.cells:
                        _set_cell_bg(c, 'FEF2F2')

        _add_footer(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_moratoriums(data: dict) -> bytes:
        """Génère le rapport des moratoires en DOCX"""
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        students = data.get('students', [])
        _add_header(doc, "Rapport — Moratoires",
                    f"Année {data['school_year']} — {len(students)} moratoire(s)")

        if students:
            tbl = doc.add_table(rows=1 + len(students), cols=5)
            tbl.style = 'Table Grid'
            _styled_table_header(tbl, ['#', 'Nom', 'Prénom', 'Classe', 'Solde (FCFA)'])
            for i, s in enumerate(students):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(str(i + 1)).font.size = Pt(8)
                row.cells[1].paragraphs[0].add_run(s['lastName']).font.size = Pt(8)
                row.cells[2].paragraphs[0].add_run(s['firstName']).font.size = Pt(8)
                row.cells[3].paragraphs[0].add_run(s['className']).font.size = Pt(8)
                r = row.cells[4].paragraphs[0].add_run(f"{s['balance']:,.0f}")
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = COLOR_RED
                if i % 2 == 0:
                    for c in row.cells:
                        _set_cell_bg(c, 'FFF7ED')
        else:
            doc.add_paragraph('Aucun moratoire enregistré pour cette période.')

        _add_footer(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def generate_payments_by_class(data: dict) -> bytes:
        """Génère le rapport des paiements par classe en DOCX"""
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        _add_header(doc, f"Paiements par Classe",
                    f"Classe : {data['class_name']} — Année {data['school_year']}")

        # Résumé financier
        summ = doc.add_table(rows=1, cols=3)
        summ.style = 'Table Grid'
        summary_data = [
            ('Total facturé', f"{data['total_invoiced']:,.0f} FCFA", '0B4B83'),
            ('Total payé', f"{data['total_paid']:,.0f} FCFA", '059669'),
            ('Reste à payer', f"{data['total_invoiced'] - data['total_paid']:,.0f} FCFA", 'DC2626'),
        ]
        for i, (label, val, color) in enumerate(summary_data):
            c = summ.cell(0, i)
            _set_cell_bg(c, color)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = COLOR_WHITE
            p2 = c.add_paragraph(label)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.runs[0] if p2.runs else p2.add_run(label)
            r2.font.size = Pt(8)
            r2.font.color.rgb = RGBColor(209, 213, 219)

        doc.add_paragraph()
        students = data.get('students', [])
        if students:
            tbl = doc.add_table(rows=1 + len(students), cols=5)
            tbl.style = 'Table Grid'
            _styled_table_header(tbl, ['Nom', 'Prénom', 'Matricule', 'Nb Paiements', 'Solde (FCFA)'])
            for i, s in enumerate(students):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(s['lastName']).font.size = Pt(8)
                row.cells[1].paragraphs[0].add_run(s['firstName']).font.size = Pt(8)
                row.cells[2].paragraphs[0].add_run(s['matricule']).font.size = Pt(8)
                row.cells[3].paragraphs[0].add_run(str(len(s.get('payments', [])))).font.size = Pt(8)
                bal = s['balance']
                r = row.cells[4].paragraphs[0].add_run(f"{bal:,.0f}")
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = COLOR_GREEN if bal <= 0 else COLOR_RED
                if i % 2 == 0:
                    for c in row.cells:
                        _set_cell_bg(c, 'F9FAFB')

        _add_footer(doc)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
